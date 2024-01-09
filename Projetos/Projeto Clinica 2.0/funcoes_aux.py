import os
import base64
import sqlite3
import calendar
from PIL import Image
from io import BytesIO
import locale
from datetime import datetime
from dash import html, dcc
import dash_bootstrap_components as dbc

def converte_imagem(imagem, nome_paciente):
    if imagem is not None:
        img_ = base64.b64decode(imagem.split(',')[1])
        img_buffer = BytesIO(img_)
        img = Image.open(img_buffer)
        img_path = os.path.join('imagens', f'{nome_paciente}.jpg')
        img.save(img_path, format='JPEG')
        return img_path
    return None

def carrega_nomes():
    conn = sqlite3.connect('base.db')
    cursor = conn.cursor()
    cursor.execute("SELECT id, Nome FROM Pacientes")
    resultados = cursor.fetchall()
    conn.close()
    nomes = {id: nome for id, nome in resultados}
    return nomes

def valida_mes_dia(mes, ano):
    if mes in [1, 3, 5, 7, 8, 10, 12]:
        return [{'label': str(dia), 'value': dia} for dia in range(1, 32)]
    elif mes == 2:
        if calendar.isleap(int(ano)):
            return [{'label': str(dia), 'value': dia} for dia in range(1, 30)]
        else:
            return [{'label': str(dia), 'value': dia} for dia in range(1, 29)]
    else:
        return [{'label': str(dia), 'value': dia} for dia in range(1, 31)]

def obter_dia_da_semana(ano, mes, dia):
    meses_dict = {'Janeiro': 1, 'Fevereiro': 2, 'Março': 3, 'Abril': 4, 'Maio': 5, 'Junho': 6,
                  'Julho': 7, 'Agosto': 8, 'Setembro': 9, 'Outubro': 10, 'Novembro': 11, 'Dezembro': 12}
    dias_semana_dict = {'Monday': 'segunda-feira', 'Tuesday': 'terca-feira', 'Wednesday': 'quarta-feira',
                        'Thursday': 'quinta-feira', 'Friday': 'sexta-feira', 'Saturday': 'sabado', 'Sunday': 'domingo'}
    numero_mes = meses_dict.get(mes)

    if numero_mes is not None:
        data = datetime(ano, numero_mes, dia)
        dia_semana = data.strftime('%A')
        return dias_semana_dict.get(dia_semana, 'Dia da semana inválido')
    else:
        pass

def relacao_consultas_dashboard(dia_da_semana):
    from app import df
    consultas_dia = df[df['Dia da Semana'] == dia_da_semana]
    return consultas_dia.to_dict('records')

def gera_modal_dinamicamente(modal_id):
    return dbc.Modal([
        dbc.ModalHeader(""),
        dbc.ModalBody(html.Div(id=f"{modal_id}-content")),
    ], id=modal_id, size="lg")

import docx
from docx.shared import Pt

def paragrafo(doc, text, style=None):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    if style:
        run.font.size = Pt(style.get('font_size', 12))
        run.font.bold = style.get('bold', False)
        run.font.italic = style.get('italic', False)

def gera_contrato_psicoterapia_01(paciente_selecionado, profissional, duracao,
                                  valor, conta, cidade, dia, mes, ano):
    doc = docx.Document()
    paragrafo(doc, 'CONTRATO DE PSICOTERAPIA INDIVIDUAL ONLINE', {'font_size': 14, 'bold': True, 'alignment': 'center'})
    paragrafo(doc, '')
    paragrafo(doc, f"Este documento visa oficializar o vínculo de tratamento de {paciente_selecionado} com o Psicólogo / Psicanalista {profissional}.", {'alignment': 'justify'})
    paragrafo(doc, '')
    paragrafo(doc, '1. Do Atendimento:')
    paragrafo(doc, f"Cada atendimento clínico terá a duração de aproximadamente {duracao} minutos, sendo realizado em horário previamente combinado, estando o terapeuta à disposição do paciente durante este período estabelecido.", {'alignment': 'justify'})
    paragrafo(doc, "Os dias e horário dos atendimentos serão combinados com o cliente, podendo variar de acordo com as necessidades de adequação da agenda do terapeuta e demanda do cliente.", {'alignment': 'justify'})
    paragrafo(doc, '')
    paragrafo(doc, '2. Do Sigilo:')
    paragrafo(doc, 'O psicólogo respeitará o sigilo profissional a fim de proteger, por meio da confiabilidade, a intimidade das pessoas, grupos ou organizações, a que tenha acesso no exercício profissional (Código de Ética do Psicólogo, artigo 9º).', {'alignment': 'justify'})
    paragrafo(doc, '')
    paragrafo(doc, '3. Do Período de Tratamento:', {'alignment': 'justify'})
    paragrafo(doc, '"A duração do tratamento psicoterápico varia consideravelmente dependendo da pessoa e da natureza das questões a serem trabalhadas. O paciente deve ser informado sobre qualquer prospecto de encurtamento ou alongamento do tratamento.', {'alignment': 'justify'})
    paragrafo(doc, '')
    paragrafo(doc, '4. Dos Honorários: ', {'alignment': 'justify'})
    paragrafo(doc, 'O pagamento será efetuado diretamente ao psicólogo nas datas combinadas no dia da primeira entrevista. Qualquer alteração no contrato ou reajuste somente poderá acontecer com o conhecimento e consentimento entre ambas as partes. ', {'alignment': 'justify'})
    paragrafo(doc, f'O Valor acordado para cada sessão é R$ {valor} . O pagamento será feito mediante Pix para a chave: {conta} . O pagamento poderá ser realizado a cada sessão ou mensalmente com desconto a combinar.', {'alignment': 'justify'})
    paragrafo(doc, '')
    paragrafo(doc, '5. Das Alterações de Agenda: ', {'alignment': 'justify'})
    paragrafo(doc, 'As desmarcações deverão ser feitas com antecedência de até 12 horas. O terapeuta deverá ser avisado no caso de imprevistos que impeçam o comparecimento do paciente. Mudanças de horário só serão possíveis quando houver disponibilidade do terapeuta.', {'alignment': 'justify'})
    paragrafo(doc, 'Sessões em que o cliente não comparecer serão cobradas normalmente. A partir de duas faltas consecutivas, sem aviso, durante o tratamento, o atendimento será considerado interrompido e o cliente poderá perder sua vaga preferencial de horário.', {'alignment': 'justify'})
    paragrafo(doc, '')
    paragrafo(doc, '')
    paragrafo(doc, f'{cidade} , {dia} de {mes} de {ano}.', {'alignment': 'justify'})
    paragrafo(doc, '')
    paragrafo(doc, '')
    paragrafo(doc, '')
    paragrafo(doc, '________________________________________________________________________________', {'alignment': 'justify'})
    paragrafo(doc, '')
    paragrafo(doc, 'Credencial nº: __________________ . Órgão Expeditor: __________________________', {'alignment': 'justify'})
    paragrafo(doc, '')
    paragrafo(doc, '')
    doc.save(f'{paciente_selecionado}_contrato_psicoterapia.docx')

def converte_arquivo(caminho_arquivo):
    with open(caminho_arquivo, 'rb') as arquivo:
        arquivo_blob = arquivo.read()
        return sqlite3.Binary(arquivo_blob)

